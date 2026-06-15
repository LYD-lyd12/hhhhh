from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置标题样式
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    h.font.name = '微软雅黑'
    h._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.bold = True

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = str(val)
            for p in table.rows[i+1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('「AI创作，豆你好玩」')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('Toby.AI创作首届全民AI创作大赛\n活动策划方案')
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for _ in range(3):
    doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_p.add_run('主办单位：中国电信江苏公司\n文档版本：V1.0\n编制日期：2026年10月')
info_run.font.size = Pt(13)
info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===== 目录 =====
doc.add_heading('目  录', level=1)
toc_items = [
    '一、活动概述',
    '二、产品核心介绍与传播口径',
    '三、活动核心目标',
    '四、活动时间规划',
    '五、参赛入口及报名流程',
    '六、赛道设置与作品要求',
    '七、评选机制（多维度综合评分体系）',
    '八、奖项体系（完整方案）',
    '九、完整活动规则',
    '十、推广执行节奏',
    '十一、预算汇总',
    '十二、预期效果',
    '十三、风险控制与应急预案',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ===== 一、活动概述 =====
doc.add_heading('一、活动概述', level=1)
add_styled_table(
    ['项目', '内容'],
    [
        ['活动名称', '「AI创作，豆你好玩」—— Toby.AI创作首届全民AI创作大赛'],
        ['活动主题', '用A豆释放创造力，人人都是AI创作者'],
        ['主办单位', '中国电信江苏公司'],
        ['参赛对象', '全体消费者（无身份、年龄上限限制）'],
        ['活动周期', '10月预热启动 → 11月作品征集 → 12月评选颁奖'],
        ['核心目标', '以赛促用、以奖促活，快速提升Toby.AI创作平台曝光度与用户活跃率'],
        ['活动预算', '实际成本约1.9万元（对外宣传面值约3万元）'],
    ]
)
doc.add_paragraph()

# ===== 二、产品核心介绍与传播口径 =====
doc.add_heading('二、产品核心介绍与传播口径', level=1)
doc.add_paragraph('Toby.AI创作是面向全量用户推出的一站式AI创作平台。平台独创A豆统一计量单位，用户完成一次购买后，可将A豆按需分配至全平台多个AI创作场景，彻底打破传统单一AI应用的会员壁垒，实现一套Token全场景通用，使用便捷性极强。')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('一句话卖点：')
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
p.add_run('"一份A豆，五大AI能力，全场景创作自由"')
doc.add_paragraph()

add_styled_table(
    ['功能模块', '功能说明', '传播亮点'],
    [
        ['AI唱作人', '支持文本一键生成原创音乐、各类音效', '"你的第一首AI单曲"'],
        ['AI生图', '通过提示词快速生成图像，内置AI润色引导，零基础可操作', '"零基础也能当设计师"'],
        ['AI视频', '接入Seedance等一线AI模型，高效生成创意短视频', '"一句话拍一部短片"'],
        ['AI播客', '实现文字转语音，一键生成专属播客内容', '"人人都有AI主播间"'],
        ['资讯Claw', '定时采集、智能分析各类信息，通过短信推送至用户手机', '"AI帮你看新闻"'],
    ]
)
doc.add_paragraph()

# ===== 三、活动核心目标 =====
doc.add_heading('三、活动核心目标', level=1)
goals = [
    ('以赛促用：', '通过创作大赛形式，引导用户深度体验Toby.AI全功能模块，培养使用习惯'),
    ('以奖促活：', '依托多层次奖项激励，提升用户活跃度与创作频次'),
    ('业务拉动：', '拉动Toby.AI销售品及电信随心选权益业务的订购转化'),
    ('内容沉淀：', '积累优质UGC内容，形成平台内容生态'),
    ('品牌传播：', '提升Toby.AI创作平台大众认知度，建立"人人可用AI创作"的品牌心智'),
]
for title, desc in goals:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# ===== 四、活动时间规划 =====
doc.add_heading('四、活动时间规划', level=1)
add_styled_table(
    ['阶段', '时间', '核心内容'],
    [
        ['预热蓄水期', '10月1日—10月15日', '活动预告、KOL种草、营业厅物料铺设、社群预热'],
        ['报名&订购期', '10月16日—10月31日', '用户订购产品、获取参赛资格、熟悉平台功能'],
        ['创作征集期', '11月1日—11月30日', '参赛者使用Toby.AI进行创作并提交作品'],
        ['作品初审期', '12月1日—12月7日', '平台审核作品合规性，筛选入围作品进入公投池'],
        ['全民投票期', '12月8日—12月21日（14天）', '入围作品公开展示，开放全民投票+专家评审（并行）'],
        ['成绩汇总期', '12月22日—12月25日', '按权重综合计算最终得分，确定获奖名单'],
        ['获奖公示期', '12月26日—12月28日（3天）', '获奖名单公示，接受异议申诉'],
        ['颁奖及兑奖', '12月29日—12月31日', '线上颁奖直播+奖品发放'],
    ]
)
doc.add_paragraph()

# ===== 五、参赛入口及报名流程 =====
doc.add_heading('五、参赛入口及报名流程', level=1)

doc.add_heading('5.1 参赛资格获取', level=2)
doc.add_paragraph('用户通过以下任一渠道完成订购（统一均价20元/月），即自动获得参赛资格：')
add_styled_table(
    ['渠道', '路径', '说明'],
    [
        ['线下营业厅', '到店告知营业员"办理Toby.AI创作大赛" → 订购专属销售品 → 现场激活', '营业员协助完成注册，可现场体验首次创作'],
        ['电信随心选权益', '电信APP → 权益中心 → 随心选 → 选取「Toby.AI创作」权益 → 激活生效', '已订购随心选用户可直接绑定权益'],
    ]
)
doc.add_paragraph()

doc.add_heading('5.2 报名流程', level=2)
doc.add_paragraph('采用"订购即报名，创作即参赛"的零门槛机制：')
steps = [
    ('Step 1 订购激活：', '通过线下营业厅或线上随心选渠道完成订购（20元/月）'),
    ('Step 2 注册登录：', '打开Toby.AI创作平台（APP/H5/小程序），使用电信手机号一键登录'),
    ('Step 3 确认参赛：', '进入「创作大赛」活动专区，点击"我要参赛"，选择参赛赛道（可多选），阅读并同意参赛须知'),
    ('Step 4 领取礼包：', '系统自动发放参赛专属300 A豆体验包，立即开始创作'),
    ('Step 5 提交作品：', '使用Toby.AI创作作品并直接提交，作品自动计入参赛记录'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)
doc.add_paragraph()

doc.add_heading('5.3 报名激励机制', level=2)
add_styled_table(
    ['激励项', '条件', '奖励'],
    [
        ['参赛有礼', '成功确认参赛', '赠送300 A豆体验包'],
        ['早鸟奖励', '前2000名确认参赛', '额外赠送200 A豆 + 限定"先锋创作者"数字徽章'],
        ['营业厅专属', '线下营业厅报名', '营业员现场指导完成首件作品，额外赠送100 A豆'],
        ['好友同行', '邀请好友成功参赛（不限人数）', '双方各得100 A豆'],
    ]
)
doc.add_paragraph()

# ===== 六、赛道设置与作品要求 =====
doc.add_heading('六、赛道设置与作品要求', level=1)

doc.add_heading('6.1 五大赛道', level=2)
add_styled_table(
    ['赛道', '对应功能', '作品形式', '作品要求'],
    [
        ['🎵 AI唱作人赛道', 'AI唱作人', '音频（MP3/WAV）', '使用AI唱作人生成的原创音乐/音效，时长30秒—5分钟，需附创作说明（100字内）'],
        ['🎨 AI生图赛道', 'AI生图', '图像（JPG/PNG）', '使用AI生图创作的视觉作品，单件或组图（≤9张），需附创作灵感说明'],
        ['🎬 AI视频赛道', 'AI视频', '视频（MP4，≤3分钟）', '使用AI视频生成的短视频，需附作品简介（100字内）'],
        ['🎙️ AI播客赛道', 'AI播客', '音频（MP3/WAV）', '使用AI播客生成的播客内容，时长1—15分钟，主题不限'],
        ['🏆 全能创作赛道', '多模块综合', '组合提交', '综合使用≥3个功能模块的跨界创作，需附整体创作说明'],
    ]
)
doc.add_paragraph()

doc.add_heading('6.2 作品提交规范', level=2)
specs = [
    ('提交方式：', '在Toby.AI创作平台内"创作大赛"专区直接提交，作品自动关联参赛账号'),
    ('提交数量：', '单赛道不限提交数量，鼓励多创作（创作量纳入评分）'),
    ('原创声明：', '提交即视为承诺作品为本人通过Toby.AI创作平台生成，禁止搬运、抄袭'),
    ('内容合规：', '作品不得包含违法违规、低俗暴力、侵犯他人权益的内容，违规即取消参赛资格'),
]
for title, desc in specs:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)
doc.add_paragraph()

# ===== 七、评选机制 =====
doc.add_heading('七、评选机制（多维度综合评分体系）', level=1)

doc.add_heading('7.1 评分维度与权重', level=2)
add_styled_table(
    ['评分维度', '权重', '评分说明'],
    [
        ['创作活跃度', '20%', '统计参赛期间用户的A豆消耗量、作品提交总数、功能模块使用覆盖度'],
        ['作品质量', '30%', '专家评审团从创意性、完成度、艺术表现力、场景应用价值四个子维度评分'],
        ['人气投票', '30%', '全民线上投票，每人每天限投3票（同一作品限投1票），投票数据真实透明'],
        ['社交传播力', '10%', '作品被分享次数、社交平台传播数据（分享按钮点击量+外部平台引用量）'],
        ['功能探索度', '10%', '参赛期间使用Toby.AI功能模块的多样性（使用越多模块得分越高，引导用户体验全功能）'],
    ]
)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('综合得分公式：')
run.bold = True
doc.add_paragraph('综合得分 = 创作活跃度得分×20% + 作品质量得分×30% + 人气投票得分×30% + 社交传播力得分×10% + 功能探索度得分×10%')
doc.add_paragraph()

doc.add_heading('7.2 评选流程', level=2)
add_styled_table(
    ['阶段', '时间', '内容'],
    [
        ['作品征集', '11月1日—11月30日', '参赛者创作并提交作品'],
        ['作品初审', '12月1日—12月7日', '平台审核作品合规性，筛选入围作品进入公投池'],
        ['全民投票', '12月8日—12月21日（14天）', '活动页公开展示入围作品，开放全民投票通道'],
        ['专家评审', '12月8日—12月21日（与投票并行）', '评审团独立打分，提升效率'],
        ['成绩汇总', '12月22日—12月25日', '按权重综合计算最终得分，确定获奖名单'],
        ['获奖公示', '12月26日（公示期3天）', '在Toby.AI创作平台及电信官方渠道公示获奖名单'],
        ['颁奖及兑奖', '12月29日—12月31日', '线上颁奖 + 奖品发放'],
    ]
)
doc.add_paragraph()

doc.add_heading('7.3 投票规则', level=2)
vote_rules = [
    ('投票资格：', '所有注册用户均可参与投票（不限于参赛者），拉动非参赛用户关注活动'),
    ('投票限额：', '每人每天3票，同一作品每天限投1票，防止刷票'),
    ('防刷机制：', '同一设备/IP/手机号限投，异常投票行为自动检测并清除无效票数'),
    ('投票激励：', '每日完成投票可抽取A豆红包（每日奖池2000 A豆），吸引持续回访'),
]
for title, desc in vote_rules:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)
doc.add_paragraph()

# ===== 八、奖项体系 =====
doc.add_heading('八、奖项体系（完整方案）', level=1)

p = doc.add_paragraph()
run = p.add_run('设计原则：')
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
principles = [
    '以江苏电信自有资源为主奖品：话费、流量、宽带、权益会员——面值感知高，实际边际成本极低',
    'A豆作为核心激励介质：平台虚拟资源，几乎零成本，但对用户有实用价值',
    '实物奖品控制在小件、实用级别：不送手机/平板，改送与电信场景相关的小硬件',
    '总预算控制在2万元以内（对外宣传面值3万+）',
]
for pr in principles:
    doc.add_paragraph(pr, style='List Bullet')
doc.add_paragraph()

doc.add_heading('8.1 主赛道奖项（全能创作赛道）', level=2)
add_styled_table(
    ['奖项等级', '名额', '奖品内容', '面值感知', '实际成本'],
    [
        ['🥇 全能金奖', '1名', '江苏电信千兆宽带免费1年 + 5,000 A豆', '≈1,800元', '≈400元'],
        ['🥈 全能银奖', '2名', '江苏电信5G畅享套餐（129元档）免费半年 + 3,000 A豆', '≈774元', '≈200元'],
        ['🥉 全能铜奖', '3名', '江苏电信100元话费 + 2,000 A豆', '≈300元', '≈120元'],
    ]
)
doc.add_paragraph()

doc.add_heading('8.2 单项赛道奖项（4条赛道，每条独立评奖）', level=2)
add_styled_table(
    ['奖项', '名额（每赛道）', '奖品内容', '面值感知', '实际成本'],
    [
        ['最佳创意奖', '1名', '天翼云盘黄金会员1年 + 2,000 A豆', '≈200元', '≈60元'],
        ['最佳人气奖', '1名', '电信50元话费 + 1,000 A豆', '≈150元', '≈60元'],
        ['优秀创作者', '3名', '电信30元话费 + 500 A豆', '≈80元', '≈40元'],
    ]
)
doc.add_paragraph()

doc.add_heading('8.3 特别奖项', level=2)
add_styled_table(
    ['奖项', '名额', '评选条件', '奖品', '面值感知'],
    [
        ['新人创作之星', '3名', '首次使用Toby.AI的新用户中表现最优', '电信50元话费 + 1,500 A豆', '≈150元'],
        ['社交传播达人', '3名', '作品分享传播数据TOP3', '电信30元话费 + 1,000 A豆', '≈80元'],
        ['最活跃创作者', '3名', '创作量（A豆消耗量）TOP3', '2,000 A豆 + 电信视频会员季卡', '≈100元'],
    ]
)
doc.add_paragraph()

doc.add_heading('8.4 全民参与奖（普惠激励，人人可得）', level=2)
add_styled_table(
    ['奖项', '达成条件', '奖励'],
    [
        ['首作有礼', '提交第1件参赛作品', '50 A豆'],
        ['创作小能手', '累计提交≥5件作品', '100 A豆 + 电信5元话费券'],
        ['全能体验官', '体验全部5个功能模块', '200 A豆 + 限定数字徽章'],
        ['人气选手', '个人作品累计≥50票', '100 A豆'],
        ['推荐有奖', '成功邀请≥3位好友参赛', '200 A豆 + 电信10元话费券'],
    ]
)
doc.add_paragraph()

doc.add_heading('8.5 奖品话术包装建议', level=2)
doc.add_paragraph('对外宣传时，不直接说"送话费"，而是包装成"江苏电信AI创作大礼包"：')
packages = [
    '千兆宽带创作网络保障（金奖）',
    '5G极速创作流量包（银奖）',
    'A豆创作基金（全员可领）',
    '天翼云盘创作素材空间（创意奖）',
]
for pkg in packages:
    doc.add_paragraph(pkg, style='List Bullet')
doc.add_paragraph('让每个奖品都和"AI创作"场景挂钩，用户拿到的不只是话费，而是一整套创作基础设施。')
doc.add_paragraph()

# ===== 九、完整活动规则 =====
doc.add_heading('九、完整活动规则', level=1)

doc.add_heading('9.1 参赛须知', level=2)
rules1 = [
    '参赛者需使用中国大陆电信手机号注册Toby.AI创作账号',
    '参赛者需订购Toby.AI创作销售品或随心选权益（含Toby.AI创作），活动期间保持业务有效状态',
    '每人可报名多个赛道，作品数量不限',
    '参赛作品必须通过Toby.AI创作平台生成，禁止使用其他工具制作后冒充AI创作',
    '作品内容须合法合规，不得含有违法违规、色情暴力、政治敏感、侵犯第三方知识产权等内容',
]
for r in rules1:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph()

doc.add_heading('9.2 作品征集规则', level=2)
rules2 = [
    '作品征集期为11月1日00:00至11月30日23:59',
    '作品须在征集期内创作并提交，历史作品不计入参赛',
    '提交作品时需选择所属赛道，并填写作品标题及创作说明',
    '单赛道提交数量不限，鼓励创作者多产出、多尝试',
]
for r in rules2:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph()

doc.add_heading('9.3 投票规则', level=2)
rules3 = [
    '投票期为12月8日至12月21日，共14天',
    '投票面向所有注册用户开放（不限于参赛者）',
    '每人每天限投3票，对同一作品每天限投1票',
    '投票通道于12月21日23:59统一关闭，逾期不补',
    '平台将对投票数据进行反作弊监测，清除机器刷票、多账号刷票等无效票',
]
for r in rules3:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph()

doc.add_heading('9.4 获奖公示及兑奖流程', level=2)
rules4 = [
    '获奖名单于12月26日在Toby.AI创作平台公告栏及电信官方公众号同步公示',
    '公示期为3天（12月26日—28日），期间接受异议申诉',
    '获奖用户须在公示期结束后7天内（至次年1月1日）通过平台完成兑奖确认，逾期视为自动放弃',
    '实物奖品将在兑奖确认后15个工作日内寄出',
    'A豆及话费奖励将在兑奖确认后3个工作日内充入获奖者账户',
]
for r in rules4:
    doc.add_paragraph(r, style='List Bullet')
doc.add_paragraph()

doc.add_heading('9.5 免责声明', level=2)
disclaimers = [
    '主办方保留对活动规则的解释权与调整权',
    '因不可抗力导致活动无法正常进行的，主办方有权暂停或终止活动',
    '参赛者提交的作品，主办方有权在活动宣传中无偿使用',
]
for d in disclaimers:
    doc.add_paragraph(d, style='List Bullet')
doc.add_paragraph()

# ===== 十、推广执行节奏 =====
doc.add_heading('十、推广执行节奏', level=1)
add_styled_table(
    ['阶段', '时间', '核心动作'],
    [
        ['预热蓄水', '10月1日—15日', '营业厅物料铺设（海报/易拉宝）、电信APP弹窗预告、短信精准触达目标用户、社交媒体悬念海报'],
        ['集中爆发', '10月16日—31日', '正式开放报名、KOL种草视频发布、营业厅"AI体验角"启动、社群裂变拉新'],
        ['深度运营', '11月全月', '每日创作打卡挑战、每周主题创作赛（周赛小奖激励）、创作者社群答疑+技巧分享'],
        ['高潮冲刺', '12月1日—21日', '投票开启、拉票裂变、社交平台作品传播、投票抽奖红包'],
        ['收官传播', '12月26日—31日', '获奖名单公示、颁奖直播、获奖作品展、活动复盘传播'],
    ]
)
doc.add_paragraph()

doc.add_heading('10.1 推广渠道', level=2)
add_styled_table(
    ['渠道', '推广方式'],
    [
        ['线下营业厅', '海报、易拉宝、"AI创作体验角"、营业员话术引导'],
        ['电信APP/短信', '弹窗推送、精准短信触达目标用户'],
        ['社交媒体', '抖音/小红书KOL种草、用户作品二次传播'],
        ['社群运营', '建立参赛者微信社群，每日创作打卡+技巧分享'],
        ['裂变机制', '邀请好友报名双方各得100 A豆'],
    ]
)
doc.add_paragraph()

# ===== 十一、预算汇总 =====
doc.add_heading('十一、预算汇总', level=1)
add_styled_table(
    ['类别', '明细', '面值感知', '实际成本估算'],
    [
        ['主赛道奖品', '金奖×1 + 银奖×2 + 铜奖×3', '≈4,400元', '≈1,000元'],
        ['单项赛道奖品（4赛道）', '每赛道：创意×1 + 人气×1 + 优秀×3', '≈6,080元', '≈1,920元'],
        ['特别奖项', '新人×3 + 传播×3 + 活跃×3', '≈1,530元', '≈600元'],
        ['全民参与奖', '预估覆盖1,500人次', '≈3,000元', '≈800元'],
        ['投票红包', '每日抽奖（14天）', '≈2,000元', '≈2,000元'],
        ['评审/运营费用', '3—5位专家/KOL评审费', '≈5,000元', '≈5,000元'],
        ['物料及推广', '营业厅物料+线上广告投放', '≈8,000元', '≈8,000元'],
        ['合计', '', '≈30,000元', '≈19,300元'],
    ]
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('成本说明：')
run.bold = True
p.add_run('奖品中电信自有资源（宽带、流量、话费）边际成本远低于面值，A豆为平台虚拟资源，实际支出极低。对外宣传可按奖品面值总额传播（约3万元），提升活动感知价值。')
doc.add_paragraph()

# ===== 十二、预期效果 =====
doc.add_heading('十二、预期效果', level=1)
add_styled_table(
    ['核心指标', '目标值'],
    [
        ['参赛报名人数', '≥ 3,000人'],
        ['活动期间作品总提交量', '≥ 10,000件'],
        ['活动页累计UV', '≥ 100,000'],
        ['投票参与人次', '≥ 30,000'],
        ['新增订购用户（含随心选）', '≥ 2,000人'],
        ['用户30日留存率', '≥ 35%'],
        ['社交媒体曝光量', '≥ 200,000次'],
    ]
)
doc.add_paragraph()

# ===== 十三、风险控制与应急预案 =====
doc.add_heading('十三、风险控制与应急预案', level=1)
add_styled_table(
    ['风险类型', '风险描述', '应对措施'],
    [
        ['参赛人数不足', '报名人数低于预期', '延长报名期1周；加大营业厅引导力度；降低参赛门槛（提交1件作品即可）'],
        ['作品质量偏低', '用户对AI创作不熟悉', '社群每日推送创作教程；设立"最佳进步奖"激励；营业员现场指导'],
        ['刷票行为', '机器/多账号刷票影响公平', '技术侧反作弊监测；异常数据自动清除；设置举报通道'],
        ['业务退订', '用户参赛后退订业务', '参赛资格与业务状态绑定；退订即取消参赛资格；设置最低在订时长'],
        ['内容合规风险', '作品包含违规内容', 'AI初审+人工复审双重审核；违规即下架并取消资格；建立内容审核标准'],
        ['技术故障', '平台宕机/提交通道异常', '提前压测；准备备用服务器；设立故障补偿机制（延长提交时间）'],
    ]
)
doc.add_paragraph()

# 保存文档
output_path = r'c:\Users\73459\Desktop\xinban\curso\Toby.AI创作全民AI创作大赛策划方案.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
